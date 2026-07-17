"""Real Word comment (批注) writing for python-docx documents.

python-docx has no public API for comments, so we manipulate the OPC
package directly: create/obtain the ``word/comments.xml`` part, append
``<w:comment>`` elements, and wrap the target paragraph with
``w:commentRangeStart`` / ``w:commentRangeEnd`` / ``w:commentReference``
markers in the main document body.
"""
from __future__ import annotations

from datetime import datetime, timezone
from typing import Iterable, Optional

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from lxml import etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
RT_COMMENTS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
CT_COMMENTS = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"

DEFAULT_AUTHOR = "Hy3 Docx Assistant"


def _get_or_create_comments_part(doc: Document):
    """Return the comments part of ``doc``, creating it (and its
    relationship + content-type override) if it does not yet exist."""
    main_part = doc.part
    comments_part = None
    for rel in main_part.rels.values():
        if rel.reltype == RT_COMMENTS:
            comments_part = rel.target_part
            break

    if comments_part is None:
        comments_elm = parse_xml(f'<w:comments xmlns:w="{W}"/>')
        blob = etree.tostring(
            comments_elm, xml_declaration=True, encoding="UTF-8", standalone=True
        )
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI

        partname = PackURI("/word/comments.xml")
        comments_part = Part(partname, CT_COMMENTS, blob, main_part.package)
        main_part.relate_to(comments_part, RT_COMMENTS)
        # base Part serializes from `_blob`; the element we created above is what we mutate
    else:
        # existing part is an XmlPart; its live `element` is what gets serialized on save
        comments_elm = comments_part.element
    return comments_part, comments_elm


def _find_paragraph(doc: Document, quote: str):
    """Find the paragraph containing ``quote`` (exact, then case-insensitive)."""
    q = (quote or "").strip()
    if not q:
        return None
    for para in doc.paragraphs:
        if q in para.text:
            return para
    ql = q.lower()
    for para in doc.paragraphs:
        if ql in para.text.lower():
            return para
    return None


def _comment_range_start(cid: int):
    el = OxmlElement("w:commentRangeStart")
    el.set(qn("w:id"), str(cid))
    return el


def _comment_range_end(cid: int):
    el = OxmlElement("w:commentRangeEnd")
    el.set(qn("w:id"), str(cid))
    return el


def _comment_reference_run(cid: int):
    r = OxmlElement("w:r")
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(cid))
    r.append(ref)
    return r


def _comment_element(cid: int, author: str, date: str, body: str):
    c = OxmlElement("w:comment")
    c.set(qn("w:id"), str(cid))
    c.set(qn("w:author"), author)
    c.set(qn("w:date"), date)
    c.set(qn("w:initials"), (author[:1] or "A").upper())
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = body
    r.append(t)
    p.append(r)
    c.append(p)
    return c


def add_comments(
    doc: Document, comments: Iterable[dict], author: str = DEFAULT_AUTHOR
) -> int:
    """Write ``comments`` (each ``{"quote": ..., "comment": ...}``) as real
    Word comments into ``doc``. Returns the number of comments written.

    Each comment is anchored to the paragraph that contains its ``quote``
    (the whole paragraph is highlighted). If the quote cannot be located, a
    new paragraph containing the quote is appended and commented instead.
    """
    comments = list(comments)
    if not comments:
        return 0

    comments_part, comments_elm = _get_or_create_comments_part(doc)
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    existing_ids = [
        int(c.get(qn("w:id")))
        for c in comments_elm.findall(qn("w:comment"))
    ]
    next_id = (max(existing_ids) + 1) if existing_ids else 0

    written = 0
    for item in comments:
        quote = (item.get("quote") or item.get("target") or item.get("text") or "").strip()
        body = (item.get("comment") or item.get("suggestion") or item.get("issue") or "").strip()
        if not body:
            continue
        para = _find_paragraph(doc, quote)
        if para is None:
            para = doc.add_paragraph(quote if quote else body)
        cid = next_id
        next_id += 1
        p = para._p
        p.insert(0, _comment_range_start(cid))
        p.append(_comment_range_end(cid))
        p.append(_comment_reference_run(cid))
        comment_body = (f"【针对】{quote}\n" if quote else "") + body
        comments_elm.append(_comment_element(cid, author, now, comment_body))
        written += 1

    if not hasattr(comments_part, "_element"):
        # base Part: serialize our mutated element into the blob
        comments_part._blob = etree.tostring(
            comments_elm, xml_declaration=True, encoding="UTF-8", standalone=True
        )
    # XmlPart (e.g. CommentsPart) mutates its live element in place; nothing else needed
    return written
