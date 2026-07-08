from __future__ import annotations

import json
from pathlib import Path
from typing import Optional

from docx import Document
from fastmcp import FastMCP
from pydantic import Field

from .comments import add_comments
from .hy3_client import Hy3Client, format_paragraphs, truncate_text

mcp = FastMCP(
    "hy3-docx-assistant",
    instructions=(
        "Hy3-powered Word document assistant. It can read a local .docx file, "
        "rewrite its paragraphs according to an instruction, or generate review comments."
    ),
)


def _require_docx_path(path: str) -> Path:
    docx_path = Path(path).expanduser().resolve()
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
    if docx_path.suffix.lower() != ".docx":
        raise ValueError(f"Only .docx files are supported: {docx_path}")
    return docx_path


def _default_output_path(input_path: Path, suffix: str) -> Path:
    return input_path.with_name(f"{input_path.stem}.{suffix}{input_path.suffix}")


def _read_docx_text(path: Path) -> str:
    doc = Document(str(path))
    return format_paragraphs(paragraph.text for paragraph in doc.paragraphs)


def _copy_with_rewritten_paragraphs(input_path: Path, output_path: Path, rewritten_paragraphs: list[str]) -> int:
    doc = Document(str(input_path))
    editable = [paragraph for paragraph in doc.paragraphs if paragraph.text.strip()]
    count = min(len(editable), len(rewritten_paragraphs))
    for paragraph, new_text in zip(editable[:count], rewritten_paragraphs[:count]):
        paragraph.text = new_text.strip()
    doc.save(str(output_path))
    return count


def _parse_json_string_array(raw: str) -> list[str]:
    text = raw.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        if lines and lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines).strip()
    start = text.find("[")
    end = text.rfind("]")
    if start != -1 and end != -1 and end > start:
        text = text[start : end + 1]
    parsed = json.loads(text)
    if not isinstance(parsed, list) or not all(isinstance(item, str) for item in parsed):
        raise ValueError("Hy3 response must be a JSON array of strings")
    return parsed


def _parse_comment_list(raw: str) -> list[dict]:
    """Parse the model response into a list of {"quote", "comment"} dicts.

    Accepts a JSON array of objects (quote/comment) or strings (comment only).
    """
    text = raw.strip()
    if text.startswith("```"):
        lines = text.splitlines()
        if lines and lines[0].startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].startswith("```"):
            lines = lines[:-1]
        text = "\n".join(lines).strip()
    start = text.find("[")
    end = text.rfind("]")
    if start != -1 and end != -1 and end > start:
        text = text[start : end + 1]
    data = json.loads(text)
    if not isinstance(data, list):
        raise ValueError("Hy3 response must be a JSON array of comment objects")
    comments = []
    for item in data:
        if isinstance(item, dict):
            comments.append(
                {
                    "quote": str(item.get("quote") or item.get("target") or item.get("text") or ""),
                    "comment": str(item.get("comment") or item.get("suggestion") or item.get("issue") or ""),
                }
            )
        elif isinstance(item, str):
            comments.append({"quote": "", "comment": item})
    return comments


@mcp.tool()
def read_docx(
    path: str = Field(description="Local path to the .docx file to read."),
    question: Optional[str] = Field(default=None, description="Optional question to answer about the document."),
) -> str:
    """Read a Word document and use Hy3 to summarize it or answer a question."""
    docx_path = _require_docx_path(path)
    text = truncate_text(_read_docx_text(docx_path))
    prompt = (
        "Document content:\n"
        f"{text}\n\n"
        "Task:\n"
        f"{question or 'Summarize this document in Chinese with key points, structure, and actionable suggestions.'}"
    )
    result = Hy3Client().chat(
        "You are a careful document reading assistant. Base your answer only on the provided document content.",
        prompt,
    )
    return f"File: {docx_path}\n\n{result}"


@mcp.tool()
def edit_docx_by_instruction(
    path: str = Field(description="Local path to the source .docx file."),
    instruction: str = Field(description="Natural-language editing instruction, such as polishing, rewriting, or translating."),
    output_path: Optional[str] = Field(default=None, description="Optional output .docx path. Defaults to <name>.hy3-edited.docx."),
) -> str:
    """Rewrite non-empty paragraphs in a Word document according to an instruction using Hy3."""
    docx_path = _require_docx_path(path)
    target_path = Path(output_path).expanduser().resolve() if output_path else _default_output_path(docx_path, "hy3-edited")
    text = truncate_text(_read_docx_text(docx_path))
    response = Hy3Client().chat(
        (
            "You rewrite Word document paragraphs. Return strict JSON only: "
            "an array of strings. Keep the same paragraph order. Do not add markdown fences."
        ),
        (
            f"Instruction:\n{instruction}\n\n"
            f"Document paragraphs:\n{text}\n\n"
            "Return rewritten paragraphs as a JSON array of strings."
        ),
    )
    try:
        rewritten = _parse_json_string_array(response)
    except (json.JSONDecodeError, ValueError) as exc:
        raise ValueError(f"Hy3 did not return valid JSON paragraphs: {exc}\nRaw response:\n{response}") from exc

    edited_count = _copy_with_rewritten_paragraphs(docx_path, target_path, rewritten)
    return (
        "Edited document generated successfully.\n"
        f"Source: {docx_path}\n"
        f"Output: {target_path}\n"
        f"Rewritten paragraphs: {edited_count}"
    )


@mcp.tool()
def suggest_docx_comments(
    path: str = Field(description="Local path to the .docx file to review."),
    comment_count: int = Field(default=5, description="Number of review comments to generate."),
    focus: Optional[str] = Field(default=None, description="Optional review focus, such as logic, style, grammar, or academic rigor."),
    output_path: Optional[str] = Field(default=None, description="Optional output .docx path with real comments written in. Defaults to <name>.comments.docx."),
) -> str:
    """Use Hy3 to generate review comments and write them as real Word comments (批注) into a new .docx file."""
    docx_path = _require_docx_path(path)
    target_path = Path(output_path).expanduser().resolve() if output_path else _default_output_path(docx_path, "comments")
    text = truncate_text(_read_docx_text(docx_path))
    response = Hy3Client().chat(
        (
            "You are a professional document reviewer. Return strict JSON only: an array of objects. "
            "Each object must have 'quote' (an exact short phrase copied verbatim from the document to anchor the comment) "
            "and 'comment' (your specific review comment in Chinese). Do not add markdown fences."
        ),
        (
            f"Document content:\n{text}\n\n"
            f"Generate {comment_count} review comments. "
            f"Review focus: {focus or 'overall structure, clarity, correctness, and style'}.\n"
            "Each comment must quote a real phrase from the document. Return a JSON array of objects."
        ),
    )
    try:
        comments = _parse_comment_list(response)
    except (json.JSONDecodeError, ValueError) as exc:
        raise ValueError(f"Hy3 did not return valid JSON comments: {exc}\nRaw response:\n{response}") from exc

    doc = Document(str(docx_path))
    written = add_comments(doc, comments, author="DeepSeek Reviewer")
    doc.save(str(target_path))
    summary = "\n".join(
        f"{i + 1}. 【{c.get('quote', '')}】 {c.get('comment', '')}"
        for i, c in enumerate(comments)
    )
    return (
        "Review comments written as real Word comments.\n"
        f"Source: {docx_path}\n"
        f"Output: {target_path}\n"
        f"Comments written: {written}\n\n"
        "--- Comments (also embedded in the output .docx) ---\n"
        f"{summary}"
    )


def main() -> None:
    mcp.run(transport="stdio")


if __name__ == "__main__":
    main()