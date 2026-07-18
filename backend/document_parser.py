"""
Document parser - extract plain text from 14 supported formats.

Each parser returns (text, meta). PDF uses pdfplumber (best quality) with
pypdf as fallback. Office/structured formats use their dedicated readers.
Everything else falls back to reading the raw bytes as UTF-8 text.
"""
import csv as csv_module
import json
import logging
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Optional, Tuple

import config

logger = logging.getLogger(__name__)


@dataclass
class ParsedDocument:
    text: str
    meta: dict


def _read_pdf(filepath: str) -> str:
    """Extract text from a PDF.

    pypdf is the stable primary parser (no native segfaults). pdfplumber is
    only used as a last-resort fallback when pypdf yields almost no text,
    because pdfplumber can segfault on certain PDFs in-process.
    """
    # 1) pypdf - stable, preferred.
    try:
        from pypdf import PdfReader
        text = "\n\n".join(
            (page.extract_text() or "") for page in PdfReader(str(filepath)).pages
        )
        if len(text.strip()) > 50:
            return text
    except Exception as e:
        logger.debug("pypdf failed: %s", e)

    # 2) pdfplumber - better for complex layouts, but used only when
    #    pypdf produced little usable text.
    try:
        import pdfplumber
        with pdfplumber.open(str(filepath)) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
        return "\n\n".join(pages)
    except Exception as e:
        logger.debug("pdfplumber failed: %s", e)

    raise RuntimeError("Could not parse PDF")


def _read_docx(filepath: str) -> str:
    from docx import Document
    doc = Document(str(filepath))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_markdown(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_json(path: Path) -> str:
    data = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
    return json.dumps(data, ensure_ascii=False, indent=2)


def _read_csv(path: Path) -> str:
    lines = []
    with path.open(encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv_module.reader(f)
        for row in reader:
            lines.append("\t".join(row))
    return "\n".join(lines)


def _read_xml(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_html(path: Path) -> str:
    from html.parser import HTMLParser

    class _TextExtractor(HTMLParser):
        def __init__(self):
            super().__init__()
            self.buf = []

        def handle_data(self, data):
            self.buf.append(data)

    text = path.read_text(encoding="utf-8", errors="ignore")
    p = _TextExtractor()
    p.feed(text)
    return "\n".join(p.buf)


def _read_rst(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_code(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


_PARSERS = {
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".txt": _read_text,
    ".md": _read_markdown,
    ".json": _read_json,
    ".csv": _read_csv,
    ".xml": _read_xml,
    ".html": _read_html,
    ".htm": _read_html,
    ".rst": _read_rst,
    ".py": _read_code,
    ".js": _read_code,
    ".ts": _read_code,
    ".doc": None,  # legacy .doc needs antiword/textract; fall back to raw
}


def parse_document(filepath) -> Tuple[str, dict]:
    """Parse `filepath` into (text, meta). Raises on unsupported format."""
    path = Path(filepath)
    ext = path.suffix.lower()
    if ext not in config.SUPPORTED_FORMATS and ext not in _PARSERS:
        raise ValueError(f"Unsupported format: {ext}")

    file_type = config.SUPPORTED_FORMATS.get(ext, ext.lstrip(".").upper())

    if ext == ".doc":
        # Legacy Word format - try to read as text best-effort.
        try:
            text = _read_text(path)
        except Exception:
            raise RuntimeError("Legacy .doc requires antiword/textract; not available")
    else:
        parser = _PARSERS.get(ext, _read_text)
        if parser is None:
            raise RuntimeError(f"No parser for {ext}")
        text = parser(path)

    meta = {
        "filename": path.name,
        "file_type": file_type,
        "size_bytes": path.stat().st_size,
        "page_count": text.count("\n\n") + 1 if ext == ".pdf" else None,
    }
    logger.info("Parsed %s (%d chars, %s)", path.name, len(text), file_type)
    return text, meta
