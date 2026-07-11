#!/usr/bin/env python3
"""Generate sample PDF and DOCX documents for document_summary / document_visualize."""

from pathlib import Path
from docx import Document
from docx.shared import Inches
from fpdf import FPDF

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "sample_data"
OUT_DIR.mkdir(exist_ok=True)

TITLE = "Q1-Q4 Sales Report"
PARAGRAPHS = [
    "This report summarizes the sales performance of the North America region across four quarters.",
    "Revenue grew steadily throughout the year, with the strongest quarter being Q4 during the holiday season.",
    "Key highlights:",
]
BULLETS = [
    "Total annual revenue reached $525,000.",
    "Unit sales increased by 21% compared to the previous year.",
    "Q4 revenue of $152,000 set a new quarterly record.",
]
TABLE = [
    ["Quarter", "Revenue ($)", "Units Sold"],
    ["Q1", "120,000", "3,400"],
    ["Q2", "135,000", "3,800"],
    ["Q3", "118,000", "3,200"],
    ["Q4", "152,000", "4,100"],
]


def generate_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading(TITLE, level=1)
    for p in PARAGRAPHS:
        doc.add_paragraph(p)
    for bullet in BULLETS:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Quarterly Performance", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(TABLE[0]):
        hdr_cells[i].text = header
    for row in TABLE[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    doc.add_paragraph("Use this document with hy3_document_summary or hy3_document_visualize.")
    doc.save(path)


class _PDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "B", 16)
        self.cell(0, 10, TITLE, new_x="LMARGIN", new_y="NEXT", align="C")
        self.ln(4)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")


def generate_pdf(path: Path) -> None:
    pdf = _PDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "", 12)

    for p in PARAGRAPHS:
        pdf.multi_cell(0, 8, p)
        pdf.ln(2)

    for bullet in BULLETS:
        pdf.set_x(18)
        pdf.multi_cell(0, 8, f"- {bullet}")

    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Quarterly Performance", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "B", 11)

    col_widths = [40, 50, 50]
    row_height = 10
    for row in TABLE:
        for i, value in enumerate(row):
            pdf.cell(col_widths[i], row_height, value, border=1)
        pdf.ln(row_height)

    pdf.ln(6)
    pdf.set_font("Helvetica", "", 12)
    pdf.multi_cell(0, 8, "Use this document with hy3_document_summary or hy3_document_visualize.")

    pdf.output(path)


if __name__ == "__main__":
    docx_path = OUT_DIR / "report.docx"
    pdf_path = OUT_DIR / "report.pdf"
    generate_docx(docx_path)
    generate_pdf(pdf_path)
    print(f"Generated {docx_path}")
    print(f"Generated {pdf_path}")
